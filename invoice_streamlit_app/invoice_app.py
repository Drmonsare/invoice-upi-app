import streamlit as st
from datetime import datetime
import random
import qrcode
from io import BytesIO
from docx import Document
from docx.shared import Inches

def format_date(date_str):
    try:
        if date_str.upper() == "NA":
            return "–"
        if '/' in date_str:
            if len(date_str.split('/')[-1]) == 2:
                date_str = date_str[:-2] + '20' + date_str[-2:]
            return datetime.strptime(date_str, "%d/%m/%Y").strftime("%d/%m/%Y")
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    except:
        return '–'

st.title("📄 Invoice Generator + UPI QR")

# User Inputs
name = st.text_input("Customer Name")
phone = st.text_input("Mobile Number (10 digits)")
month = st.selectbox("Month", ["January","February","March","April","May","June","July","August","September","October","November","December"])
year = st.text_input("Year", value="2025")
date = st.date_input("Invoice Date", value=datetime.today())
upi_id = st.text_input("Your UPI ID (e.g. yourname@upi)")

# Items
st.subheader("Enter Items")
items = []
for i in range(5):  # limit to 5 items
    with st.expander(f"Item {i+1}"):
        desc = st.text_input(f"Description {i+1}", key=f"desc{i}")
        amount = st.number_input(f"Amount ₹", min_value=0.0, step=0.5, key=f"amt{i}")
        item_date = st.text_input(f"Date of item (dd/mm/yy or NA)", value="NA", key=f"date{i}")
        if desc and amount:
            items.append((format_date(item_date), desc, amount))

if st.button("Generate Invoice"):
    if len(phone) != 10 or not phone.isdigit():
        st.error("📵 Mobile number must be 10 digits.")
    elif not upi_id or not name:
        st.error("❗ Name and UPI ID are required.")
    else:
        total = sum(x[2] for x in items)
        inv_no = random.randint(250000, 259999)

        doc = Document()
        doc.add_heading('INVOICE', level=1)
        doc.add_paragraph(f"STATEMENT\n{month}/{year}")
        doc.add_paragraph(f"DATE: {date.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"INVOICE #: {inv_no}")
        doc.add_paragraph(f"Customer: {name}")
        doc.add_paragraph(f"Mobile: {phone}")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DATE & DESCRIPTION'
        hdr_cells[1].text = 'AMOUNT ₹'

        for d, desc, amt in items:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{d} - {desc}"
            row_cells[1].text = f"{amt:.2f}"

        doc.add_paragraph(f"\nTOTAL: ₹ {total:.2f}")

        # QR Code
        upi_qr = f"upi://pay?pa={upi_id}&pn={name}&am={total:.2f}&cu=INR"
        qr_img = qrcode.make(upi_qr)
        qr_io = BytesIO()
        qr_img.save(qr_io)
        qr_io.seek(0)

        with open("upi_qr.png", "wb") as f:
            f.write(qr_io.getbuffer())
        doc.add_picture("upi_qr.png", width=Inches(2.0))

        # Save and Download
        filename = f"Invoice_{inv_no}.docx"
        doc.save(filename)
        with open(filename, "rb") as f:
            st.download_button("⬇ Download Invoice", data=f, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")